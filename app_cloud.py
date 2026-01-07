import streamlit as st
import os
import time
from io import BytesIO
from docx import Document
from datetime import datetime
from ollama import Client

from pubmed_client import PubMedClient
from analyzer import TrendAnalyzer
from main import sanitize_topic

# ========================================
# Streamlit Secrets Configuration
# ========================================
# HOW TO SET SECRETS IN STREAMLIT CLOUD:
# 1. Go to your app dashboard on share.streamlit.io
# 2. Click "Settings" → "Secrets"
# 3. Add the following in TOML format:
#
# ENTREZ_EMAIL = "your_email@example.com"
# ENTREZ_API_KEY = "your_ncbi_api_key_here"
# OLLAMA_API_KEY = "your_ollama_cloud_api_key"
# OLLAMA_MODEL = "gptoss-120b:cloud"
#
# Get your Ollama API key from: https://ollama.com/settings/keys
# ========================================

# Page Config
st.set_page_config(
    page_title="Pharmacoepi Trend Generator",
    page_icon="💊",
    layout="wide"
)

def get_secret(key: str, default: str = "") -> str:
    """Safely retrieve secrets with fallback to environment variables (for local testing)."""
    try:
        return st.secrets.get(key, os.getenv(key, default))
    except Exception:
        return os.getenv(key, default)

def generate_docx(result, topic, date_str):
    """Generates a DOCX file in memory with TTE format."""
    doc = Document()
    doc.add_heading(f'Pharmacoepidemiology Research Trend Report: {topic}', 0)
    doc.add_paragraph(f'Date: {date_str}')
    
    doc.add_heading('1. Research Landscape', level=1)
    doc.add_paragraph(result.landscape_summary)
    
    doc.add_heading('2. Evidence Gaps (Feasibility Filtered)', level=1)
    for gap in result.identified_gaps:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"[{gap.category}] ").bold = True
        p.add_run(f"{gap.gap} (Feasibility: {gap.feasibility_status})")
        
    doc.add_heading('3. Research Hypotheses (Target Trial Emulation)', level=1)
    for i, hyp in enumerate(result.hypotheses, 1):
        doc.add_heading(f'Hypothesis {i}: {hyp.title}', level=2)
        
        # Research Question
        p = doc.add_paragraph()
        p.add_run('Research Question (PICO): ').bold = True
        p.add_run(hyp.research_question)
        
        # Rationale
        p = doc.add_paragraph()
        p.add_run('Rationale: ').bold = True
        p.add_run(hyp.rationale)
        
        # Study Design
        p = doc.add_paragraph()
        p.add_run('Study Design: ').bold = True
        p.add_run(hyp.study_design)
        
        # Target Trial Components
        doc.add_heading('Target Trial Components', level=3)
        ttc = hyp.target_trial_components
        doc.add_paragraph(f"Population: {ttc.population}")
        doc.add_paragraph(f"Intervention: {ttc.intervention}")
        doc.add_paragraph(f"Comparator: {ttc.comparator}")
        doc.add_paragraph(f"Outcome: {ttc.outcome_operational_def}")
        doc.add_paragraph(f"Follow-up: {ttc.follow_up}")
        doc.add_paragraph(f"Time Zero: {ttc.time_zero_definition}")
        
        # Bias Mitigation
        doc.add_heading('Bias Mitigation', level=3)
        bm = hyp.bias_mitigation
        doc.add_paragraph(f"Key Confounders: {', '.join(bm.key_confounders)}")
        doc.add_paragraph(f"Negative Control Outcome: {bm.negative_control_outcome}")
        doc.add_paragraph(f"Sensitivity Analysis: {bm.sensitivity_analysis or 'E-value'}")
        
        # Feasibility
        doc.add_heading('Feasibility Assessment', level=3)
        fa = hyp.feasibility_assessment
        doc.add_paragraph(f"Data Source: {fa.data_source_suitability}")
        doc.add_paragraph(f"Sample Size: {fa.expected_sample_size}")
        doc.add_paragraph(f"Challenges: {fa.potential_challenges}")
        
        doc.add_paragraph("_" * 50)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_markdown_string(result, topic, date_str):
    """Generates Markdown content with TTE format."""
    md = f"# Pharmacoepidemiology Research Trend Report: {topic}\n"
    md += f"**Date:** {date_str}\n\n"
    
    md += "## 1. Research Landscape\n"
    md += f"{result.landscape_summary}\n\n"
    
    md += "## 2. Evidence Gaps (Feasibility Filtered)\n"
    for gap in result.identified_gaps:
        md += f"- **[{gap.category}]** {gap.gap} (Feasibility: {gap.feasibility_status})\n"
    md += "\n"
    
    md += "## 3. Research Hypotheses (Target Trial Emulation)\n"
    for i, hyp in enumerate(result.hypotheses, 1):
        md += f"### Hypothesis {i}: {hyp.title}\n\n"
        md += f"**Research Question (PICO):** {hyp.research_question}\n\n"
        md += f"**Rationale:** {hyp.rationale}\n\n"
        md += f"**Study Design:** {hyp.study_design}\n\n"
        
        md += "#### Target Trial Components\n"
        ttc = hyp.target_trial_components
        md += f"| Component | Definition |\n|---|---|\n"
        md += f"| Population | {ttc.population} |\n"
        md += f"| Intervention | {ttc.intervention} |\n"
        md += f"| Comparator | {ttc.comparator} |\n"
        md += f"| Outcome | {ttc.outcome_operational_def} |\n"
        md += f"| Follow-up | {ttc.follow_up} |\n"
        md += f"| Time Zero | {ttc.time_zero_definition} |\n\n"
        
        md += "#### Bias Mitigation\n"
        bm = hyp.bias_mitigation
        md += f"- **Key Confounders:** {', '.join(bm.key_confounders)}\n"
        md += f"- **Negative Control:** {bm.negative_control_outcome}\n"
        md += f"- **Sensitivity:** {bm.sensitivity_analysis or 'E-value'}\n\n"
        
        md += "#### Feasibility\n"
        fa = hyp.feasibility_assessment
        md += f"- **Data Source:** {fa.data_source_suitability}\n"
        md += f"- **Sample Size:** {fa.expected_sample_size}\n"
        md += f"- **Challenges:** {fa.potential_challenges}\n\n"
        
        md += "---\n"
    return md

def main():
    st.title("💊 Pharmacoepidemiology Research Trend Generator")
    st.markdown("Generate research hypotheses based on real-time PubMed trends.")

    # Sidebar Controls
    with st.sidebar:
        st.header("⚙️ Config")
        
        # Get credentials from secrets
        default_model = get_secret("OLLAMA_MODEL", "gptoss-120b:cloud")
        model_name = st.text_input("Model", value=default_model)
        
        max_results = st.slider("Abstracts", 5, 50, 20)
        
        # Check credentials
        api_key = get_secret("ENTREZ_API_KEY", "")
        email = get_secret("ENTREZ_EMAIL", "")
        ollama_api_key = get_secret("OLLAMA_API_KEY", "")
        
        st.subheader("📋 Credentials Status")
        
        # NCBI Status
        if api_key and email:
            st.success("✅ NCBI Credentials Configured")
        else:
            st.error("❌ Missing NCBI Credentials")
            st.caption("Add `ENTREZ_EMAIL` and `ENTREZ_API_KEY` to Secrets")
        
        # Ollama Cloud Status
        if ollama_api_key:
            st.success("✅ Ollama Cloud API Key Configured")
        else:
            st.error("❌ Missing Ollama Cloud API Key")
            st.caption("Add `OLLAMA_API_KEY` to Secrets")
            st.caption("Get key from: [ollama.com/settings/keys](https://ollama.com/settings/keys)")
        
        st.divider()
        st.caption("💡 **Quick Links:**")
        st.caption("• [NCBI API Key](https://www.ncbi.nlm.nih.gov/account/)")
        st.caption("• [Ollama Cloud API Key](https://ollama.com/settings/keys)")
        
    topic_input = st.text_input("Enter Topic", placeholder="e.g., GLP-1 Agonists")
    
    # Main Action
    if st.button("🚀 Analyze Trends & Generate Hypotheses", type="primary"):
        if not topic_input:
            st.error("Please enter a research topic.")
            return

        # Check credentials before proceeding
        if not email:
            st.error("⚠️ ENTREZ_EMAIL not configured. Please add it to Streamlit Secrets.")
            st.stop()
        
        if not ollama_api_key:
            st.error("⚠️ OLLAMA_API_KEY not configured. Please add it to Streamlit Secrets.")
            st.info("Get your API key from: https://ollama.com/settings/keys")
            st.stop()

        safe_topic = sanitize_topic(topic_input)
        date_str = datetime.now().strftime('%Y-%m-%d')
        
        # Initialize Ollama Cloud Client
        try:
            ollama_client = Client(
                host='https://ollama.com',
                headers={'Authorization': f'Bearer {ollama_api_key}'}
            )
            st.info("🔗 Ollama Cloud client initialized")
        except Exception as e:
            st.error(f"⚠️ Failed to initialize Ollama Cloud client: {e}")
            st.stop()
        
        # Initialize PubMed Client
        os.environ["ENTREZ_EMAIL"] = email
        if api_key:
            os.environ["ENTREZ_API_KEY"] = api_key
        
        client = PubMedClient()
        analyzer = TrendAnalyzer(model_name=model_name, ollama_client=ollama_client)
        
        result = None
        
        # Status container
        with st.status("Processing...", expanded=True) as status:
            # 1. Fetch
            status.write(f"🔍 Searching PubMed for '{topic_input}'...")
            query = f"{topic_input} AND (pharmacoepidemiology OR 'real world evidence' OR observational)"
            
            try:
                abstracts = client.fetch_abstracts(query, max_results=max_results)
            except Exception as e:
                status.update(label="PubMed Fetch Failed", state="error")
                st.error(f"Failed to fetch from PubMed: {e}")
                import traceback
                st.code(traceback.format_exc())
                st.stop()
            
            if not abstracts:
                status.update(label="No papers found.", state="error")
                st.error("No abstracts found for this topic. Try a broader term.")
                st.stop()
            
            status.write(f"✅ Found {len(abstracts)} papers. Analyzing with Ollama Cloud...")
            
            # 2. Analyze
            try:
                result = analyzer.analyze_and_hypothesize(abstracts, topic_input)
                
                # DEBUG: Check result validity
                if result:
                    st.write(f"🔍 DEBUG - Result type: {type(result)}")
                    st.write(f"🔍 DEBUG - Trend summary length: {len(result.trend_summary)}")
                    st.write(f"🔍 DEBUG - Gaps count: {len(result.identified_gaps)}")
                    st.write(f"🔍 DEBUG - Hypotheses count: {len(result.hypotheses)}")
                    status.update(label="Analysis Complete! 🎉", state="complete", expanded=False)
                else:
                    st.error("⚠️ Analysis returned None!")
                    st.stop()
                    
            except Exception as e:
                status.update(label="Analysis Failed", state="error")
                st.error(f"❌ Error during analysis: {e}")
                st.error("**Possible issues:**")
                st.error("• Invalid Ollama API key")
                st.error("• Model not available on Ollama Cloud")
                st.error("• Network/connection error")
                st.error("• API rate limit exceeded")
                
                # Show full traceback for debugging
                import traceback
                with st.expander("🐛 Full Error Traceback"):
                    st.code(traceback.format_exc())
                st.stop()

        # Display Results
        if result:
            st.success("✅ Results loaded successfully!")
            
            tab1, tab2, tab3 = st.tabs(["📋 Report", "🔬 Hypotheses", "📑 Data"])
            
            with tab1:
                st.subheader("📊 Research Landscape")
                st.info(result.landscape_summary)
                
                st.subheader("🔍 Evidence Gaps (Feasibility Filtered)")
                for gap in result.identified_gaps:
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.warning(f"**{gap.category}**: {gap.gap}")
                    with col2:
                        st.caption(f"Feasibility: {gap.feasibility_status}")
                    
            with tab2:
                for i, hyp in enumerate(result.hypotheses, 1):
                    with st.expander(f"🧪 {i}. {hyp.title}", expanded=(i==1)):
                        # Research Question
                        st.markdown(f"**📝 Research Question (PICO):**")
                        st.info(hyp.research_question)
                        
                        # Rationale
                        st.markdown(f"**💡 Rationale:**")
                        st.write(hyp.rationale)
                        
                        st.divider()
                        
                        # Target Trial Components
                        st.markdown("### 🎯 Target Trial Emulation Components")
                        ttc = hyp.target_trial_components
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            st.markdown("**Population (Eligibility):**")
                            st.caption(ttc.population)
                            st.markdown("**Intervention:**")
                            st.caption(ttc.intervention)
                            st.markdown("**Comparator:**")
                            st.caption(ttc.comparator)
                        with col2:
                            st.markdown("**Outcome (Operational Def):**")
                            st.caption(ttc.outcome_operational_def)
                            st.markdown("**Follow-up:**")
                            st.caption(ttc.follow_up)
                            st.markdown("**Time Zero:**")
                            st.caption(ttc.time_zero_definition)
                        
                        st.divider()
                        
                        # Bias Mitigation
                        st.markdown("### 🛡️ Bias Mitigation")
                        bm = hyp.bias_mitigation
                        
                        st.markdown("**Key Confounders for PS:**")
                        st.write(", ".join(bm.key_confounders) if bm.key_confounders else "Not specified")
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            st.markdown("**Negative Control Outcome:**")
                            st.caption(bm.negative_control_outcome)
                        with col2:
                            st.markdown("**Sensitivity Analysis:**")
                            st.caption(bm.sensitivity_analysis or "E-value calculation")
                        
                        st.divider()
                        
                        # Feasibility
                        st.markdown("### 📊 Feasibility Assessment")
                        fa = hyp.feasibility_assessment
                        
                        st.markdown("**Data Source Suitability:**")
                        st.write(fa.data_source_suitability)
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            st.markdown("**Expected Sample Size:**")
                            st.caption(fa.expected_sample_size)
                        with col2:
                            st.markdown("**Potential Challenges:**")
                            st.caption(fa.potential_challenges)
            
            with tab3:
                st.caption(f"**Total Papers Analyzed**: {len(abstracts)}")
                for p in abstracts: 
                    st.text(f"{p.get('date', 'N/A')} | {p['title']}")
            
            # Download Section
            st.divider()
            st.subheader("📥 Download Report")
            
            col_d1, col_d2 = st.columns(2)
            
            filename_base = f"report_{safe_topic}_{datetime.now().strftime('%Y%m%d')}"
            
            # Markdown Download
            with col_d1:
                st.download_button(
                    label="📄 Download Markdown",
                    data=generate_markdown_string(result, topic_input, date_str),
                    file_name=f"{filename_base}.md",
                    mime="text/markdown"
                )
            
            # DOCX Download
            docx_buffer = generate_docx(result, topic_input, date_str)
            with col_d2:
                st.download_button(
                    label="📝 Download Word (.docx)",
                    data=docx_buffer,
                    file_name=f"{filename_base}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error("⚠️ No results to display. Result is None.")

if __name__ == "__main__":
    main()
