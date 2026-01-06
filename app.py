import streamlit as st
import os
import time
from io import BytesIO
from docx import Document
from datetime import datetime

from pubmed_client import PubMedClient
from analyzer import TrendAnalyzer
from main import sanitize_topic

# ========================================
# Streamlit Secrets Configuration
# ========================================
# HOW TO SET SECRETS IN STREAMLIT CLOUD:
# 1. Go to your app dashboard on share.streamlit.io
# 2. Click "Settings" → "Secrets"
# 3. Add the following in TOML format:
#
# ENTREZ_EMAIL = "your_email@example.com"
# ENTREZ_API_KEY = "your_ncbi_api_key_here"
# OLLAMA_MODEL = "gptoss-120b:cloud"
# OLLAMA_HOST = "https://your-ollama-api-url.com"
#
# ========================================

# Page Config
st.set_page_config(
    page_title="Pharmacoepi Trend Generator",
    page_icon="💊",
    layout="wide"
)

def get_secret(key: str, default: str = "") -> str:
    """Safely retrieve secrets with fallback to environment variables (for local testing)."""
    try:
        return st.secrets.get(key, os.getenv(key, default))
    except Exception:
        return os.getenv(key, default)

def generate_docx(result, topic, date_str):
    """Generates a DOCX file in memory."""
    doc = Document()
    doc.add_heading(f'Pharmacoepidemiology Research Trend Report: {topic}', 0)
    doc.add_paragraph(f'Date: {date_str}')
    
    doc.add_heading('1. Current Research Trends', level=1)
    doc.add_paragraph(result.trend_summary)
    
    doc.add_heading('2. Identified Evidence Gaps', level=1)
    for gap in result.identified_gaps:
        doc.add_paragraph(gap, style='List Bullet')
        
    doc.add_heading('3. Proposed Research Hypotheses', level=1)
    for i, hyp in enumerate(result.hypotheses, 1):
        doc.add_heading(f'Hypothesis {i}: {hyp.title}', level=2)
        p = doc.add_paragraph()
        p.add_run('Rationale: ').bold = True
        p.add_run(hyp.rationale)
        
        p2 = doc.add_paragraph()
        p2.add_run('Suggested Methodology: ').bold = True
        p2.add_run(hyp.methodology)
        
        doc.add_paragraph("_" * 50)

    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_markdown_string(result, topic, date_str):
    """Generates the Markdown content as a string."""
    md = f"# Pharmacoepidemiology Research Trend Report: {topic}\n"
    md += f"**Date:** {date_str}\n\n"
    
    md += "## 1. Current Research Trends\n"
    md += f"{result.trend_summary}\n\n"
    
    md += "## 2. Identified Evidence Gaps\n"
    for gap in result.identified_gaps:
            md += f"- {gap}\n"
    md += "\n"
    
    md += "## 3. Proposed Research Hypotheses\n"
    for i, hyp in enumerate(result.hypotheses, 1):
        md += f"### Hypothesis {i}: {hyp.title}\n"
        md += f"**Rationale:** {hyp.rationale}\n\n"
        md += f"**Suggested Methodology:** {hyp.methodology}\n\n"
        md += "---\n"
    return md

def main():
    st.title("💊 Pharmacoepidemiology Research Trend Generator")
    st.markdown("Generate research hypotheses based on real-time PubMed trends.")

    # Sidebar Controls
    with st.sidebar:
        st.header("⚙️ Config")
        
        # Get model name from secrets or use default
        default_model = get_secret("OLLAMA_MODEL", "gptoss-120b:cloud")
        model_name = st.text_input("Model", value=default_model)
        
        max_results = st.slider("Abstracts", 5, 50, 20)
        
        # Check if API key is configured
        api_key = get_secret("ENTREZ_API_KEY", "")
        email = get_secret("ENTREZ_EMAIL", "")
        
        if api_key and email:
            st.success("✅ NCBI Credentials Configured")
        else:
            st.error("❌ Missing NCBI Credentials")
            st.info("👉 Add `ENTREZ_EMAIL` and `ENTREZ_API_KEY` to Secrets")
        
        st.divider()
        st.caption("💡 **Configuration Tips:**")
        st.caption("• Set secrets in Streamlit Cloud dashboard")
        st.caption("• NCBI API key: [Get here](https://www.ncbi.nlm.nih.gov/account/)")
        
    topic_input = st.text_input("Enter Topic", placeholder="e.g., GLP-1 Agonists")
    
    # Main Action
    if st.button("🚀 Analyze Trends & Generate Hypotheses", type="primary"):
        if not topic_input:
            st.error("Please enter a research topic.")
            return

        # Check credentials before proceeding
        if not email:
            st.error("⚠️ ENTREZ_EMAIL not configured. Please add it to Streamlit Secrets.")
            st.stop()

        safe_topic = sanitize_topic(topic_input)
        date_str = datetime.now().strftime('%Y-%m-%d')
        
        # Initialize Logic
        # Pass email and API key explicitly to override environment variables
        os.environ["ENTREZ_EMAIL"] = email
        if api_key:
            os.environ["ENTREZ_API_KEY"] = api_key
        
        client = PubMedClient()
        analyzer = TrendAnalyzer(model_name=model_name)
        
        result = None
        
        # Status container
        with st.status("Processing...", expanded=True) as status:
            # 1. Fetch
            status.write(f"🔍 Searching PubMed for '{topic_input}'...")
            query = f"{topic_input} AND (pharmacoepidemiology OR 'real world evidence' OR observational)"
            abstracts = client.fetch_abstracts(query, max_results=max_results)
            
            if not abstracts:
                status.update(label="No papers found.", state="error")
                st.error("No abstracts found for this topic. Try a broader term.")
                st.stop()
            
            status.write(f"✅ Found {len(abstracts)} papers. Analyzing...")
            
            # 2. Analyze
            try:
                result = analyzer.analyze_and_hypothesize(abstracts, topic_input)
                status.update(label="Analysis Complete! 🎉", state="complete", expanded=False)
            except Exception as e:
                status.update(label="Analysis Failed", state="error")
                st.error(f"Error during analysis: {e}")
                st.stop()

        # Display Results
        if result:
            tab1, tab2, tab3 = st.tabs(["Report", "Hypotheses", "Data"])
            with tab1:
                st.subheader("📊 Executive Summary")
                st.info(result.trend_summary)
                
                st.subheader("🔍 Evidence Gaps")
                for gap in result.identified_gaps:
                    st.warning(f"- {gap}")
                    
            with tab2:
                for i, hyp in enumerate(result.hypotheses, 1):
                    with st.expander(f"{i}. {hyp.title}", expanded=True):
                        st.write(f"**Rationale:** {hyp.rationale}")
                        st.info(f"**Methodology:** {hyp.methodology}")
            
            with tab3:
                for p in abstracts: st.text(f"{p['date']} | {p['title']}")
            
            # Download Section
            st.divider()
            st.subheader("📥 Download Report")
            
            col_d1, col_d2 = st.columns(2)
            
            filename_base = f"report_{safe_topic}_{datetime.now().strftime('%Y%m%d')}"
            
            # Markdown Download
            with col_d1:
                st.download_button(
                    label="📄 Download Markdown",
                    data=generate_markdown_string(result, topic_input, date_str),
                    file_name=f"{filename_base}.md",
                    mime="text/markdown"
                )
            
            # DOCX Download
            docx_buffer = generate_docx(result, topic_input, date_str)
            with col_d2:
                st.download_button(
                    label="📝 Download Word (.docx)",
                    data=docx_buffer,
                    file_name=f"{filename_base}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
